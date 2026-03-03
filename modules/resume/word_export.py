# -*- coding: utf-8 -*-
"""
ATS-friendly Word resume exporter for Indian job market.

Rules:
- Single column, no tables for layout (only content tables if absolutely needed)
- Standard fonts only (Calibri)
- No text boxes, no graphics, no headers/footers with content
- Clean bullet points via python-docx List Bullet style
- Naukri RChilli and LinkedIn parser compatible
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


DARK  = RGBColor(0x1F, 0x27, 0x3A)
BLUE  = RGBColor(0x1B, 0x4F, 0x9C)
GREY  = RGBColor(0x55, 0x55, 0x55)
LINE  = RGBColor(0xCC, 0xCC, 0xCC)


# == Helpers ==================================================================

def _hr(doc: Document) -> None:
    """Add a thin horizontal rule."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE
    run.font.name = "Calibri"
    _hr(doc)


def _body_run(para, text: str, bold: bool = False, size: int = 10, color: RGBColor = None) -> None:
    run = para.add_run(text)
    run.bold = bool(bold)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def _bullet(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    para.paragraph_format.left_indent  = Inches(0.2)
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK


# == Main export ==============================================================

def export_to_docx(resume_data: dict) -> bytes:
    """
    Build an ATS-friendly Word document from Claude's resume JSON.
    Returns bytes suitable for st.download_button().
    """
    doc = Document()

    # Page margins — narrow for more content space
    section = doc.sections[0]
    section.page_width       = Inches(8.5)
    section.page_height      = Inches(11)
    section.left_margin      = Inches(0.75)
    section.right_margin     = Inches(0.75)
    section.top_margin       = Inches(0.6)
    section.bottom_margin    = Inches(0.6)

    # Default paragraph spacing
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # == HEADER: Name + contact ===============================================
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(resume_data.get("name", "").upper())
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.name = "Calibri"
    name_run.font.color.rgb = DARK

    target_para = doc.add_paragraph()
    target_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target_para.paragraph_format.space_after = Pt(2)
    _body_run(target_para, resume_data.get("target_title", ""), bold=True, size=11, color=BLUE)

    contact_parts = []
    if resume_data.get("phone"):   contact_parts.append(resume_data["phone"])
    if resume_data.get("email"):   contact_parts.append(resume_data["email"])
    if resume_data.get("location"): contact_parts.append(resume_data["location"])
    if resume_data.get("linkedin"): contact_parts.append(resume_data["linkedin"])

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(4)
    _body_run(contact_para, "  |  ".join(contact_parts), size=9, color=GREY)

    _hr(doc)

    # == PROFESSIONAL SUMMARY =================================================
    _section_heading(doc, "Professional Summary")
    summary_para = doc.add_paragraph()
    summary_para.paragraph_format.space_after = Pt(4)
    _body_run(summary_para, resume_data.get("summary", ""), size=10)

    # == KEY SKILLS ===========================================================
    skills = resume_data.get("skills", {})
    if skills:
        _section_heading(doc, "Key Skills")
        all_skills = []
        for category, skill_list in skills.items():
            if skill_list:
                all_skills.extend(skill_list)

        # Render skills in two visual rows (plain text, ATS-safe)
        mid = (len(all_skills) + 1) // 2
        col1 = all_skills[:mid]
        col2 = all_skills[mid:]
        max_len = max(len(col1), len(col2))
        for i in range(max_len):
            s1 = f"• {col1[i]}" if i < len(col1) else ""
            s2 = f"• {col2[i]}" if i < len(col2) else ""
            line = doc.add_paragraph()
            line.paragraph_format.space_before = Pt(1)
            line.paragraph_format.space_after  = Pt(1)
            _body_run(line, f"{s1:<38}{s2}", size=10)

    # == WORK EXPERIENCE ======================================================
    experience = resume_data.get("experience", [])
    if experience:
        _section_heading(doc, "Work Experience")
        for job in experience:
            # Company + Period line
            header_para = doc.add_paragraph()
            header_para.paragraph_format.space_before = Pt(6)
            header_para.paragraph_format.space_after  = Pt(1)
            _body_run(header_para, job.get("company", ""), bold=True, size=11, color=DARK)
            period = job.get("period", "")
            location = job.get("location", "")
            right_text = f"  |  {period}" + (f"  |  {location}" if location else "")
            _body_run(header_para, right_text, bold=False, size=10, color=GREY)

            # Job title
            title_para = doc.add_paragraph()
            title_para.paragraph_format.space_after = Pt(2)
            _body_run(title_para, job.get("title", ""), bold=True, size=10, color=BLUE)

            # Bullets
            for bullet in job.get("bullets", []):
                _bullet(doc, bullet)

    # == EDUCATION ============================================================
    education = resume_data.get("education", [])
    if education:
        _section_heading(doc, "Education")
        for edu in education:
            edu_para = doc.add_paragraph()
            edu_para.paragraph_format.space_before = Pt(4)
            edu_para.paragraph_format.space_after  = Pt(1)
            _body_run(edu_para, edu.get("degree", ""), bold=True, size=10, color=DARK)
            inst = edu.get("institution", "")
            year = edu.get("year", "")
            grade = edu.get("grade", "")
            sub_parts = [p for p in [inst, year, grade] if p]
            sub_para = doc.add_paragraph()
            sub_para.paragraph_format.space_after = Pt(2)
            _body_run(sub_para, "  |  ".join(sub_parts), size=9.5, color=GREY)

    # == CERTIFICATIONS =======================================================
    certs = resume_data.get("certifications", [])
    if certs:
        _section_heading(doc, "Certifications")
        for cert in certs:
            _bullet(doc, cert)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def get_filename(resume_data: dict) -> str:
    name   = resume_data.get("name", "Resume").replace(" ", "_")
    role   = resume_data.get("target_title", "").replace(" ", "_").replace("/", "-")
    return f"{name}_{role}_Resume.docx"
